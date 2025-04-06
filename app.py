# # # app.py

# import streamlit as st
# from document_processor import extract_text_from_pdf, extract_text_from_docx, clean_text
# from analyzer import analyze_rfp
# from report_generator import format_analysis_report
# from utils import display_markdown_report, highlight_verdict
# #from pdf_generator import generate_rfp_pdf 
# import fitz  # PyMuPDF
# from docx import Document
# import google.generativeai as genai
# # Page Configuration
# st.set_page_config(
#     page_title="AutoRFP.AI - Government RFP Analyzer",
#     layout="wide",
#     page_icon="📄",
# )

# st.markdown("""
#     <h1 style='text-align: center;'>📑 AutoRFP.AI</h1>
#     <h4 style='text-align: center;'>Automating RFP Analysis with Gemini + RAG + Agentic Workflows ⚙️</h4>
#     <hr>
# """, unsafe_allow_html=True)

# # Upload Section
# with st.sidebar:
#     st.header("📂 Document + Gemini")

#     # 1. API Key
#     GEMINI_API_KEY = "AIzaSyAxNo1eAvrfi2g6FbXNQLDRcW3bdf293SI"  # or make it editable

#     # 2. File Upload
#     uploaded_file = st.file_uploader("📤 Upload PDF or DOCX", type=["pdf", "docx"])

#     # 3. Reset Chat
#     if st.button("🔄 Reset Chat"):
#         st.session_state.chat_history = []
#         st.rerun()

#     # 4. Gemini Setup
#     if GEMINI_API_KEY:
#         try:
#             genai.configure(api_key=GEMINI_API_KEY)
#             model = genai.GenerativeModel("models/gemini-1.5-pro")
#         except Exception as e:
#             st.error(f"Gemini Init Error: {str(e)}")
#             st.stop()
#     else:
#         st.warning("Please add your Gemini API Key")
#         st.stop()

#     # 5. File Text Extraction
#     def extract_text_from_pdf(pdf_file) -> str:
#         try:
#             text = ""
#             with fitz.open(stream=pdf_file.read(), filetype="pdf") as doc:
#                 for page in doc:
#                     text += page.get_text()
#             return text.strip()
#         except Exception as e:
#             return f"❌ PDF Error: {str(e)}"

#     def extract_text_from_docx(docx_file) -> str:
#         try:
#             doc = Document(docx_file)
#             return "\n".join([para.text for para in doc.paragraphs])
#         except Exception as e:
#             return f"❌ DOCX Error: {str(e)}"

#     def clean_text(text: str) -> str:
#         return " ".join(text.replace("\n", " ").replace("\r", " ").split())

#     file_text = ""
#     if uploaded_file:
#         if uploaded_file.name.endswith(".pdf"):
#             file_text = extract_text_from_pdf(uploaded_file)
#         else:
#             file_text = extract_text_from_docx(uploaded_file)

#         file_text = clean_text(file_text)

#         if file_text.startswith("❌"):
#             st.error(file_text)
#         else:
#             st.success("✅ Document loaded. Ask your question below.")

#     # 6. Chat Section
#     if "chat_history" not in st.session_state:
#         st.session_state.chat_history = []

#     if file_text:
#         user_question = st.text_input("🧠 Ask a question:")

#         if user_question:
#             prompt = f"""
# You are a document analysis expert. Answer clearly and accurately based on the following content.

# --- DOCUMENT START ---
# {file_text[:12000]}
# --- DOCUMENT END ---

# --- QUESTION ---
# {user_question}
# """
#             try:
#                 response = model.generate_content(prompt)
#                 answer = response.text
#             except Exception as e:
#                 answer = f"❌ Gemini API Error: {str(e)}"

#             st.session_state.chat_history.append((user_question, answer))

#     if st.session_state.chat_history:
#         st.subheader("🗨️ Conversation")
#         for q, a in st.session_state.chat_history:
#             st.markdown(f"**User:** {q}")
#             st.markdown(f"**Bot:** {a}")

# col1, col2 = st.columns(2)

# with col1:
#     uploaded_rfp = st.file_uploader("📄 Upload RFP (PDF)", type=["pdf"])

# with col2:
#     uploaded_docx = st.file_uploader("🏢 Upload Company Profile (DOCX)", type=["docx"])

# if uploaded_rfp and uploaded_docx:
#     # Process files
#     with st.expander("📜 Preview Extracted RFP Text"):
#         raw_text = extract_text_from_pdf(uploaded_rfp)
#         rfp_text = clean_text(raw_text)
#         st.text_area("Raw RFP Extracted Text", rfp_text[:3000], height=300)

#     with st.expander("🏢 Preview Company Profile"):
#         company_profile = extract_text_from_docx(uploaded_docx)
#         st.text_area("Extracted Company Profile", company_profile[:3000], height=200)

#     if st.button("🚀 Run Full Analysis"):
#         with st.spinner("Analyzing the RFP using Gemini Agents..."):
#             analysis = analyze_rfp(rfp_text, company_profile)
#             report = format_analysis_report(analysis)

#         st.success("✅ Analysis Complete")
#         highlight_verdict(analysis["Eligibility Verdict"])
#         display_markdown_report(report)


# else:
#     st.info("📌 Upload an RFP PDF and Company Profile DOCX to begin analysis.")

# # from pdf_generator import generate_rfp_pdf  # ✅ Add this at the top of app.py

# # # 📄 PDF Download Section
# # if "analysis" in locals() and analysis:
# #     if st.button("📄 Download Submission-Ready PDF"):
# #         pdf_path = generate_rfp_pdf(analysis)
# #         with open(pdf_path, "rb") as f:
# #             st.download_button(
# #                 label="📥 Download Final RFP Report PDF",
# #                 data=f,
# #                 file_name="rfp_analysis_report.pdf",
# #                 mime="application/pdf"
# #             )


# # After displaying analysis report

# #########################################################################chatbot
# # ✅ 🧠 Chatbot Section

# if st.button("🔄 Reset All"):
#     st.session_state.clear()
#     st.rerun()
import streamlit as st
import streamlit as st
from document_processor import extract_text_from_pdf, extract_text_from_docx, clean_text
from document_processor import extract_text_from_pdf, extract_text_from_docx, clean_text
from analyzer import analyze_rfp
from analyzer import analyze_rfp
from report_generator import format_analysis_report
from report_generator import format_analysis_report
from utils import display_markdown_report, highlight_verdict
from utils import display_markdown_report, highlight_verdict
from pdf_generator import generate_rfp_pdf  # ✅ PDF generation support
#from pdf_generator import generate_rfp_pdf
import fitz  # PyMuPDF
import fitz  # PyMuPDF
from docx import Document
from docx import Document
import google.generativeai as genai
import google.generativeai as genai
# Page Configuration
# Page Configuration
st.set_page_config(
st.set_page_config(
    page_title="AutoRFP.AI - Government RFP Analyzer",
    page_title="AutoRFP.AI - Government RFP Analyzer",
    layout="wide",
    layout="wide",
    page_icon="📄",
    page_icon="📄",
)
)
st.markdown("""
st.markdown("""
    <h1 style='text-align: center;'>📑 AutoRFP.AI</h1>
    <h1 style='text-align: center;'>📑 AutoRFP.AI</h1>
    <h4 style='text-align: center;'>Automating RFP Analysis with Gemini + RAG + Agentic Workflows ⚙️</h4>
    <h4 style='text-align: center;'>Automating RFP Analysis with Gemini + RAG + Agentic Workflows ⚙️</h4>
    <hr>
    <hr>
""", unsafe_allow_html=True)
""", unsafe_allow_html=True)
# Upload Section
# Upload Section
with st.sidebar:
with st.sidebar:
    st.header("📂 Document + Gemini")
    st.header("📂 Document + Gemini")
    # 1. API Key
    # 1. API Key
    GEMINI_API_KEY = "AIzaSyAxNo1eAvrfi2g6FbXNQLDRcW3bdf293SI"
    GEMINI_API_KEY = "AIzaSyAxNo1eAvrfi2g6FbXNQLDRcW3bdf293SI"  # or make it editable
    # 2. File Upload
    # 2. File Upload
    uploaded_file = st.file_uploader("📤 Upload PDF or DOCX", type=["pdf", "docx"])
    uploaded_file = st.file_uploader("📤 Upload PDF or DOCX", type=["pdf", "docx"])
    # 3. Reset Chat
    # 3. Reset Chat
    if st.button("🔄 Reset Chat"):
    if st.button("🔄 Reset Chat"):
        st.session_state.chat_history = []
        st.session_state.chat_history = []
        st.rerun()
        st.rerun()
    # 4. Gemini Setup
    # 4. Gemini Setup
    if GEMINI_API_KEY:
    if GEMINI_API_KEY:
        try:
        try:
            genai.configure(api_key=GEMINI_API_KEY)
            genai.configure(api_key=GEMINI_API_KEY)
            model = genai.GenerativeModel("models/gemini-1.5-pro")
            model = genai.GenerativeModel("models/gemini-1.5-pro")
        except Exception as e:
        except Exception as e:
            st.error(f"Gemini Init Error: {str(e)}")
            st.error(f"Gemini Init Error: {str(e)}")
            st.stop()
            st.stop()
    else:
    else:
        st.warning("Please add your Gemini API Key")
        st.warning("Please add your Gemini API Key")
        st.stop()
        st.stop()
    # 5. File Text Extraction
    # 5. File Text Extraction
    def extract_text_from_pdf(pdf_file) -> str:
    def extract_text_from_pdf(pdf_file) -> str:
        try:
        try:
            text = ""
            text = ""
            with fitz.open(stream=pdf_file.read(), filetype="pdf") as doc:
            with fitz.open(stream=pdf_file.read(), filetype="pdf") as doc:
                for page in doc:
                for page in doc:
                    text += page.get_text()
                    text += page.get_text()
            return text.strip()
            return text.strip()
        except Exception as e:
        except Exception as e:
            return f"❌ PDF Error: {str(e)}"
            return f"❌ PDF Error: {str(e)}"
    def extract_text_from_docx(docx_file) -> str:
    def extract_text_from_docx(docx_file) -> str:
        try:
        try:
            doc = Document(docx_file)
            doc = Document(docx_file)
            return "\n".join([para.text for para in doc.paragraphs])
            return "\n".join([para.text for para in doc.paragraphs])
        except Exception as e:
        except Exception as e:
            return f"❌ DOCX Error: {str(e)}"
            return f"❌ DOCX Error: {str(e)}"
    def clean_text(text: str) -> str:
    def clean_text(text: str) -> str:
        return " ".join(text.replace("\n", " ").replace("\r", " ").split())
        return " ".join(text.replace("\n", " ").replace("\r", " ").split())
    file_text = ""
    file_text = ""
    if uploaded_file:
    if uploaded_file:
        if uploaded_file.name.endswith(".pdf"):
        if uploaded_file.name.endswith(".pdf"):
            file_text = extract_text_from_pdf(uploaded_file)
            file_text = extract_text_from_pdf(uploaded_file)
        else:
        else:
            file_text = extract_text_from_docx(uploaded_file)
            file_text = extract_text_from_docx(uploaded_file)
        file_text = clean_text(file_text)
        file_text = clean_text(file_text)
        if file_text.startswith("❌"):
        if file_text.startswith("❌"):
            st.error(file_text)
            st.error(file_text)
        else:
        else:
            st.success("✅ Document loaded. Ask your question below.")
            st.success("✅ Document loaded. Ask your question below.")
    # 6. Chat Section
    # 6. Chat Section
    if "chat_history" not in st.session_state:
    if "chat_history" not in st.session_state:
        st.session_state.chat_history = []
        st.session_state.chat_history = []
    if file_text:
    if file_text:
        user_question = st.text_input("🧠 Ask a question:")
        user_question = st.text_input("🧠 Ask a question:")
        if user_question:
        if user_question:
            prompt = f"""
            prompt = f"""
You are a document analysis expert. Answer clearly and accurately based on the following content.
You are a document analysis expert. Answer clearly and accurately based on the following content.
--- DOCUMENT START ---
--- DOCUMENT START ---
{file_text[:12000]}
{file_text[:12000]}
--- DOCUMENT END ---
--- DOCUMENT END ---
--- QUESTION ---
--- QUESTION ---
{user_question}
{user_question}
"""
"""
            try:
            try:
                response = model.generate_content(prompt)
                response = model.generate_content(prompt)
                answer = response.text
                answer = response.text
            except Exception as e:
            except Exception as e:
                answer = f"❌ Gemini API Error: {str(e)}"
                answer = f"❌ Gemini API Error: {str(e)}"
            st.session_state.chat_history.append((user_question, answer))
            st.session_state.chat_history.append((user_question, answer))
    if st.session_state.chat_history:
    if st.session_state.chat_history:
        st.subheader("🗨️ Conversation")
        st.subheader("🗨️ Conversation")
        for q, a in st.session_state.chat_history:
        for q, a in st.session_state.chat_history:
            st.markdown(f"**User:** {q}")
            st.markdown(f"**User:** {q}")
            st.markdown(f"**Bot:** {a}")
            st.markdown(f"**Bot:** {a}")
col1, col2 = st.columns(2)
col1, col2 = st.columns(2)
with col1:
with col1:
    uploaded_rfp = st.file_uploader("📄 Upload RFP (PDF)", type=["pdf"])
    uploaded_rfp = st.file_uploader("📄 Upload RFP (PDF)", type=["pdf"])
with col2:
with col2:
    uploaded_docx = st.file_uploader("🏢 Upload Company Profile (DOCX)", type=["docx"])
    uploaded_docx = st.file_uploader("🏢 Upload Company Profile (DOCX)", type=["docx"])
if uploaded_rfp and uploaded_docx:
if uploaded_rfp and uploaded_docx:
    # Process files
    # Process files
    with st.expander("📜 Preview Extracted RFP Text"):
    with st.expander("📜 Preview Extracted RFP Text"):
        raw_text = extract_text_from_pdf(uploaded_rfp)
        raw_text = extract_text_from_pdf(uploaded_rfp)
        rfp_text = clean_text(raw_text)
        rfp_text = clean_text(raw_text)
        st.text_area("Raw RFP Extracted Text", rfp_text[:3000], height=300)
        st.text_area("Raw RFP Extracted Text", rfp_text[:3000], height=300)
    with st.expander("🏢 Preview Company Profile"):
    with st.expander("🏢 Preview Company Profile"):
        company_profile = extract_text_from_docx(uploaded_docx)
        company_profile = extract_text_from_docx(uploaded_docx)
        st.text_area("Extracted Company Profile", company_profile[:3000], height=200)
        st.text_area("Extracted Company Profile", company_profile[:3000], height=200)
    if st.button("🚀 Run Full Analysis"):
    if st.button("🚀 Run Full Analysis"):
        with st.spinner("Analyzing the RFP using Gemini Agents..."):
        with st.spinner("Analyzing the RFP using Gemini Agents..."):
            analysis = analyze_rfp(rfp_text, company_profile)
            analysis = analyze_rfp(rfp_text, company_profile)
            report = format_analysis_report(analysis)
            report = format_analysis_report(analysis)
            st.session_state.analysis = analysis
            st.session_state.report = report
        st.success("✅ Analysis Complete")
        st.success("✅ Analysis Complete")
        highlight_verdict(analysis["Eligibility Verdict"])
        highlight_verdict(analysis["Eligibility Verdict"])
        display_markdown_report(report)
        display_markdown_report(report)
        # ✅ PDF Download Section
        if st.button("📄 Download Submission-Ready PDF"):
            pdf_path = generate_rfp_pdf(analysis)
            with open(pdf_path, "rb") as f:
                st.download_button(
                    label="📥 Download Final RFP Report PDF",
                    data=f,
                    file_name="rfp_analysis_report.pdf",
                    mime="application/pdf"
                )
else:
else:
    st.info("📌 Upload an RFP PDF and Company Profile DOCX to begin analysis.")
    st.info("📌 Upload an RFP PDF and Company Profile DOCX to begin analysis.")
# from pdf_generator import generate_rfp_pdf  # ✅ Add this at the top of app.py
# # 📄 PDF Download Section
# if "analysis" in locals() and analysis:
#     if st.button("📄 Download Submission-Ready PDF"):
#         pdf_path = generate_rfp_pdf(analysis)
#         with open(pdf_path, "rb") as f:
#             st.download_button(
#                 label="📥 Download Final RFP Report PDF",
#                 data=f,
#                 file_name="rfp_analysis_report.pdf",
#                 mime="application/pdf"
#             )
# After displaying analysis report
#########################################################################chatbot
# ✅ 🧠 Chatbot Section
# ✅ 🧠 Chatbot Section
if st.button("🔄 Reset All"):
if st.button("🔄 Reset All"):
    st.session_state.clear()
    st.session_state.clear()
    st.rerun()
    st.rerun()
 
 		
import streamlit as st
from document_processor import extract_text_from_pdf, extract_text_from_docx, clean_text
from analyzer import analyze_rfp
from report_generator import format_analysis_report
from utils import display_markdown_report, highlight_verdict
from pdf_generator import generate_rfp_pdf  # ✅ PDF generation support
import fitz  # PyMuPDF
from docx import Document
import google.generativeai as genai

# Page Configuration
st.set_page_config(
    page_title="AutoRFP.AI - Government RFP Analyzer",
    layout="wide",
    page_icon="📄",
)

st.markdown("""
    <h1 style='text-align: center;'>📑 AutoRFP.AI</h1>
    <h4 style='text-align: center;'>Automating RFP Analysis with Gemini + RAG + Agentic Workflows ⚙️</h4>
    <hr>
""", unsafe_allow_html=True)

# Upload Section
with st.sidebar:
    st.header("📂 Document + Gemini")

    # 1. API Key
    GEMINI_API_KEY = "AIzaSyAxNo1eAvrfi2g6FbXNQLDRcW3bdf293SI"

    # 2. File Upload
    uploaded_file = st.file_uploader("📤 Upload PDF or DOCX", type=["pdf", "docx"])

    # 3. Reset Chat
    if st.button("🔄 Reset Chat"):
        st.session_state.chat_history = []
        st.rerun()

    # 4. Gemini Setup
    if GEMINI_API_KEY:
        try:
            genai.configure(api_key=GEMINI_API_KEY)
            model = genai.GenerativeModel("models/gemini-1.5-pro")
        except Exception as e:
            st.error(f"Gemini Init Error: {str(e)}")
            st.stop()
    else:
        st.warning("Please add your Gemini API Key")
        st.stop()

    # 5. File Text Extraction
    def extract_text_from_pdf(pdf_file) -> str:
        try:
            text = ""
            with fitz.open(stream=pdf_file.read(), filetype="pdf") as doc:
                for page in doc:
                    text += page.get_text()
            return text.strip()
        except Exception as e:
            return f"❌ PDF Error: {str(e)}"

    def extract_text_from_docx(docx_file) -> str:
        try:
            doc = Document(docx_file)
            return "\n".join([para.text for para in doc.paragraphs])
        except Exception as e:
            return f"❌ DOCX Error: {str(e)}"

    def clean_text(text: str) -> str:
        return " ".join(text.replace("\n", " ").replace("\r", " ").split())

    file_text = ""
    if uploaded_file:
        if uploaded_file.name.endswith(".pdf"):
            file_text = extract_text_from_pdf(uploaded_file)
        else:
            file_text = extract_text_from_docx(uploaded_file)

        file_text = clean_text(file_text)

        if file_text.startswith("❌"):
            st.error(file_text)
        else:
            st.success("✅ Document loaded. Ask your question below.")

    # 6. Chat Section
    if "chat_history" not in st.session_state:
        st.session_state.chat_history = []

    if file_text:
        user_question = st.text_input("🧠 Ask a question:")

        if user_question:
            prompt = f"""
You are a document analysis expert. Answer clearly and accurately based on the following content.

--- DOCUMENT START ---
{file_text[:12000]}
--- DOCUMENT END ---

--- QUESTION ---
{user_question}
"""
            try:
                response = model.generate_content(prompt)
                answer = response.text
            except Exception as e:
                answer = f"❌ Gemini API Error: {str(e)}"

            st.session_state.chat_history.append((user_question, answer))

    if st.session_state.chat_history:
        st.subheader("🗨️ Conversation")
        for q, a in st.session_state.chat_history:
            st.markdown(f"**User:** {q}")
            st.markdown(f"**Bot:** {a}")

col1, col2 = st.columns(2)

with col1:
    uploaded_rfp = st.file_uploader("📄 Upload RFP (PDF)", type=["pdf"])

with col2:
    uploaded_docx = st.file_uploader("🏢 Upload Company Profile (DOCX)", type=["docx"])

if uploaded_rfp and uploaded_docx:
    # Process files
    with st.expander("📜 Preview Extracted RFP Text"):
        raw_text = extract_text_from_pdf(uploaded_rfp)
        rfp_text = clean_text(raw_text)
        st.text_area("Raw RFP Extracted Text", rfp_text[:3000], height=300)

    with st.expander("🏢 Preview Company Profile"):
        company_profile = extract_text_from_docx(uploaded_docx)
        st.text_area("Extracted Company Profile", company_profile[:3000], height=200)

    if st.button("🚀 Run Full Analysis"):
        with st.spinner("Analyzing the RFP using Gemini Agents..."):
            analysis = analyze_rfp(rfp_text, company_profile)
            report = format_analysis_report(analysis)
            st.session_state.analysis = analysis
            st.session_state.report = report

        st.success("✅ Analysis Complete")
        highlight_verdict(analysis["Eligibility Verdict"])
        display_markdown_report(report)

        # ✅ PDF Download Section
        if st.button("📄 Download Submission-Ready PDF"):
            pdf_path = generate_rfp_pdf(analysis)
            with open(pdf_path, "rb") as f:
                st.download_button(
                    label="📥 Download Final RFP Report PDF",
                    data=f,
                    file_name="rfp_analysis_report.pdf",
                    mime="application/pdf"
                )
else:
    st.info("📌 Upload an RFP PDF and Company Profile DOCX to begin analysis.")

# ✅ 🧠 Chatbot Section
if st.button("🔄 Reset All"):
    st.session_state.clear()
    st.rerun()

